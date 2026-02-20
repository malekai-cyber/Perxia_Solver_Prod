"""
Generador de Documentación Técnica - Perxia Solver
Genera documentación profesional en .docx y PDF
Uso: python scripts/generate_docs.py
"""

import json
import os
import sys
from datetime import date
from pathlib import Path

# ─────────────────────────────────────────────
# Paleta de colores corporativa
# ─────────────────────────────────────────────
C_NAVY       = (13,  27,  63)    # Azul oscuro - títulos principales
C_BLUE       = (30,  90, 168)    # Azul medio  - encabezados sección
C_ACCENT     = (0,  160, 198)    # Cyan        - acento / badges
C_ORANGE     = (230, 126,  34)   # Naranja     - highlights
C_LIGHT_GRAY = (245, 247, 250)   # Fondo celdas pares
C_MID_GRAY   = (189, 195, 199)   # Bordes
C_DARK_GRAY  = (44,  62,  80)    # Texto normal
C_WHITE      = (255, 255, 255)
C_SUCCESS    = (39, 174,  96)    # Verde
C_WARNING    = (243, 156,  18)   # Amarillo


# ══════════════════════════════════════════════════════════════
#  DATOS DEL DOCUMENTO  (modificar aquí para refrescar)
# ══════════════════════════════════════════════════════════════
class DocData:
    PROJECT         = "Perxia Solver – Agente de Análisis Inteligente de Oportunidades"
    VERSION         = "v1.0.0"
    DATE            = date.today().strftime("%d de %B de %Y")
    COMPANY         = "Periferia IT Group / CBIT"
    SUBTITLE        = "Documentación Técnica de Producción"
    REPO            = "https://github.com/malekai-cyber/Perxia_Solver_Prod"
    FUNCTION_URL    = "https://func-analyzer-prod.azurewebsites.net/api/analyze"
    AZURE_RG        = "rg_perxia_solver_prod"
    AZURE_FUNC      = "func-analyzer-prod"
    AZURE_SUB       = "Azure Production"

    SUMMARY = (
        "Perxia Solver es un agente de inteligencia artificial desarrollado sobre "
        "Azure Functions que analiza automáticamente oportunidades comerciales "
        "entrantes desde Dynamics 365 CRM vía Power Automate. El agente evalúa "
        "requisitos técnicos, estima esfuerzo y riesgo, y recomienda los equipos "
        "de trabajo más adecuados consultando la base de conocimiento organizacional "
        "almacenada en Azure AI Search. El resultado se entrega como Adaptive Card "
        "en Microsoft Teams y PDF almacenado en Blob Storage."
    )

    FLOW_STEPS = [
        ("1", "Creación / actualización de Oportunidad",
         "Dynamics 365 CRM genera un trigger en Power Automate al crear o modificar "
         "una oportunidad comercial."),
        ("2", "Envío HTTP a Azure Functions",
         "Power Automate realiza un POST HTTP al endpoint de la Function App "
         "incluyendo el payload JSON con datos de la oportunidad."),
        ("3", "Validación del Payload",
         "El modelo Pydantic OpportunityPayload valida y normaliza los datos "
         "incluyendo limpieza de HTML en la descripción."),
        ("4", "Consulta a la Base de Conocimiento",
         "El SearchService consulta Azure AI Search obteniendo el catálogo completo "
         "de torres/equipos con sus competencias, líderes y contactos."),
        ("5", "Análisis con GPT-4o-mini",
         "El OpenAIService envía el contexto de la oportunidad y los equipos al "
         "modelo, obteniendo un análisis estructurado en JSON con estimaciones, "
         "riesgos, recomendaciones de equipo y próximos pasos."),
        ("6", "Enriquecimiento con datos reales",
         "El orquestador sobreescribe los datos de líderes y emails con los valores "
         "auténticos del catálogo, evitando que la IA invente información."),
        ("7", "Generación de PDF",
         "PDFGenerator crea un documento profesional con el análisis completo y lo "
         "sube a Azure Blob Storage retornando una URL SAS con vigencia de 90 días."),
        ("8", "Almacenamiento en Cosmos DB",
         "El registro completo del análisis se persiste en Azure Cosmos DB para "
         "auditoría, historial y consultas posteriores."),
        ("9", "Generación Adaptive Card",
         "Se construye la Adaptive Card JSON compatible con Microsoft Teams v1.4 "
         "con el resumen ejecutivo, equipos recomendados y enlace al PDF."),
        ("10", "Respuesta a Power Automate",
         "La Function retorna HTTP 200 con el JSON completo incluyendo análisis, "
         "card y URLs. Power Automate usa la Adaptive Card para notificar al equipo "
         "comercial en Teams."),
    ]

    SERVICES = [
        {
            "name": "OpenAIService",
            "file": "shared/services/openai_service.py",
            "description": "Motor de razonamiento del agente. Conecta con Azure OpenAI GPT-4o-mini para analizar oportunidades y generar recomendaciones estructuradas en JSON.",
            "env_vars": [
                "AZURE_OPENAI_ENDPOINT", "AZURE_OPENAI_KEY",
                "AZURE_OPENAI_DEPLOYMENT_NAME", "AZURE_OPENAI_API_VERSION"
            ],
            "key_method": "analyze_opportunity(opportunity_text, available_teams) → Dict",
        },
        {
            "name": "SearchService",
            "file": "shared/services/search_service.py",
            "description": "Gestiona la consulta a Azure AI Search para recuperar el catálogo completo de torres y equipos con sus competencias.",
            "env_vars": [
                "AZURE_SEARCH_ENDPOINT", "AZURE_SEARCH_KEY", "AZURE_SEARCH_INDEX_TEAMS"
            ],
            "key_method": "get_all_teams() → List[Dict]",
        },
        {
            "name": "BlobStorageService",
            "file": "shared/services/blob_storage_service.py",
            "description": "Gestiona el almacenamiento de PDFs generados en Azure Blob Storage y genera URLs SAS con tiempo de expiración configurable.",
            "env_vars": [
                "AZURE_STORAGE_CONNECTION_STRING", "AZURE_STORAGE_CONTAINER_NAME"
            ],
            "key_method": "upload_pdf(pdf_bytes, blob_name) → str (SAS URL)",
        },
        {
            "name": "CosmosDBService",
            "file": "shared/services/cosmos_service.py",
            "description": "Persiste los registros de análisis en Azure Cosmos DB NoSQL para historial y auditoría.",
            "env_vars": [
                "COSMOS_ENDPOINT", "COSMOS_KEY",
                "COSMOS_DATABASE_NAME", "COSMOS_CONTAINER_NAME"
            ],
            "key_method": "save_analysis(record: Dict) → Dict",
        },
    ]

    APP_SETTINGS = [
        # (Nombre, Requerido, Descripción, Ejemplo)
        ("AZURE_OPENAI_ENDPOINT",          "✅ Sí", "URL del recurso Azure OpenAI",
         "https://mi-openai.openai.azure.com/"),
        ("AZURE_OPENAI_KEY",               "✅ Sí", "Clave API de Azure OpenAI", "abc123..."),
        ("AZURE_OPENAI_DEPLOYMENT_NAME",   "✅ Sí", "Nombre del deployment del modelo", "gpt-4o-mini"),
        ("AZURE_OPENAI_API_VERSION",       "✅ Sí", "Versión de la API OpenAI", "2024-10-21"),
        ("AZURE_SEARCH_ENDPOINT",          "✅ Sí", "URL del servicio Azure AI Search",
         "https://mi-search.search.windows.net"),
        ("AZURE_SEARCH_KEY",               "✅ Sí", "Clave admin de Azure AI Search", "xyz789..."),
        ("AZURE_SEARCH_INDEX_TEAMS",       "✅ Sí", "Nombre del índice de equipos", "torres-index"),
        ("AZURE_STORAGE_CONNECTION_STRING","✅ Sí", "Cadena de conexión Blob Storage",
         "DefaultEndpointsProtocol=https;..."),
        ("AZURE_STORAGE_CONTAINER_NAME",   "✅ Sí", "Contenedor de PDFs", "analysis-pdfs"),
        ("COSMOS_ENDPOINT",                "⚠️ Opc", "URL de Cosmos DB",
         "https://mi-cosmos.documents.azure.com:443/"),
        ("COSMOS_KEY",                     "⚠️ Opc", "Clave Cosmos DB", "..."),
        ("COSMOS_DATABASE_NAME",           "⚠️ Opc", "Base de datos Cosmos DB", "perxia-db"),
        ("COSMOS_CONTAINER_NAME",          "⚠️ Opc", "Contenedor Cosmos DB", "analyses"),
    ]

    PAYLOAD_FIELDS = [
        ("opportunityid",      "string", "✅", "ID único de la oportunidad en Dynamics"),
        ("name",               "string", "✅", "Nombre / título de la oportunidad"),
        ("description",        "string", "⚠️", "Descripción técnica (acepta HTML)"),
        ("description_funcional", "string", "⚠️", "Descripción funcional limpia"),
        ("estimatedvalue",     "number", "⚠️", "Valor estimado en USD"),
        ("currencyid",         "string", "⚠️", "Moneda (ej: USD, COP)"),
        ("customerid",         "string", "⚠️", "ID del cliente en Dataverse"),
        ("customername",       "string", "⚠️", "Nombre del cliente"),
        ("statecode",          "integer","⚠️", "Estado (0=Abierta, 1=Ganada, 2=Perdida)"),
        ("createdon",          "string", "⚠️", "Fecha ISO 8601 de creación"),
        ("sdk_message",        "string", "⚠️", "Tipo de evento (Create / Update)"),
    ]

    ERROR_CODES = [
        ("VALIDATION_ERROR",      "400", "Payload inválido o campos requeridos faltantes"),
        ("IMPORT_ERROR",          "200", "Error al importar módulos internos – revisar dependencias"),
        ("SERVICE_NOT_CONFIGURED","200", "Variables de entorno de servicio no configuradas"),
        ("AI_ANALYSIS_ERROR",     "200", "El modelo AI no retornó resultado válido"),
        ("PROCESSING_ERROR",      "200", "Error interno durante el procesamiento"),
    ]

    CICD_STEPS = [
        ("Checkout", "actions/checkout@v4",
         "Descarga el código fuente del repositorio."),
        ("Setup Python 3.13", "actions/setup-python@v5",
         "Configura el intérprete Python en el runner de GitHub."),
        ("Install dependencies", "pip install -r requirements.txt",
         "Instala todas las dependencias del proyecto."),
        ("Lint – flake8", "python -m flake8 . --max-line-length=120",
         "Valida estilo de código. Falla el pipeline si hay errores."),
        ("Tests – pytest", "python -m pytest tests/ -v",
         "Ejecuta la suite de pruebas unitarias (22 tests)."),
        ("Package artifact", "zip -r release.zip ...",
         "Empaqueta AnalyzeOpportunity/, shared/, host.json y requirements.txt."),
        ("Deploy to Azure", "azure/functions-action@v1",
         "Despliega el ZIP en func-analyzer-prod usando el Publish Profile."),
    ]


# ══════════════════════════════════════════════════════════════
#  GENERADOR  .DOCX
# ══════════════════════════════════════════════════════════════
def generate_docx(output_path: str, torres: list) -> None:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # ── Configurar márgenes ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Helpers de estilo ──────────────────────────────────────
    def rgb(*c): return RGBColor(*c)

    def set_cell_bg(cell, r, g, b):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), f"{r:02X}{g:02X}{b:02X}")
        tcPr.append(shd)

    def add_heading(text, level=1, color=C_NAVY):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = rgb(*color)
        run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}[level])
        if level == 1:
            run.font.name = "Calibri"
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:space"), "4")
            bottom.set(qn("w:color"), f"{C_ACCENT[0]:02X}{C_ACCENT[1]:02X}{C_ACCENT[2]:02X}")
            pBdr.append(bottom)
            pPr.append(pBdr)
        return p

    def add_body(text, indent=False):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        if indent:
            p.paragraph_format.left_indent = Cm(0.8)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = rgb(*C_DARK_GRAY)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = rgb(*C_DARK_GRAY)
        return p

    def add_kv(key, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{key}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(value)
        r2.font.size = Pt(10)
        r2.font.color.rgb = rgb(*C_DARK_GRAY)

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.8)
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = rgb(198, 40, 40)

    def styled_table(headers, rows, col_widths=None, header_color=C_NAVY):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_bg(hdr_cells[i], *header_color)
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = rgb(*C_WHITE)
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ri, row in enumerate(rows):
            cells = table.add_row().cells
            bg = C_LIGHT_GRAY if ri % 2 == 0 else C_WHITE
            for i, val in enumerate(row):
                set_cell_bg(cells[i], *bg)
                p = cells[i].paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(9)
        if col_widths:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell.width = Cm(col_widths[j])
        doc.add_paragraph()  # espaciado tras tabla
        return table

    # ── PORTADA ──────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_title.add_run(DocData.PROJECT)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = rgb(*C_NAVY)
    r.font.name = "Calibri"

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cover_sub.add_run(DocData.SUBTITLE)
    r2.font.size = Pt(14)
    r2.font.color.rgb = rgb(*C_BLUE)

    doc.add_paragraph()
    for label, value in [
        ("Versión",  DocData.VERSION),
        ("Fecha",    DocData.DATE),
        ("Empresa",  DocData.COMPANY),
        ("Repo",     DocData.REPO),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = p.add_run(f"{label}: ")
        rb.bold = True
        rb.font.size = Pt(11)
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = rgb(*C_DARK_GRAY)

    doc.add_page_break()

    # ── 1. RESUMEN EJECUTIVO ──────────────────────────────────
    add_heading("1. Resumen Ejecutivo", 1)
    add_body(DocData.SUMMARY)

    doc.add_paragraph()
    add_kv("Function App",     DocData.AZURE_FUNC)
    add_kv("Resource Group",   DocData.AZURE_RG)
    add_kv("Endpoint",         DocData.FUNCTION_URL)
    add_kv("Modelo IA",        "GPT-4o-mini (Azure OpenAI)")
    add_kv("Runtime",          "Python 3.13 – Azure Functions v4")
    add_kv("Timeout",          "10 minutos")
    add_kv("CI/CD",            "GitHub Actions → Publish Profile")

    # ── 2. ARQUITECTURA ──────────────────────────────────────
    doc.add_paragraph()
    add_heading("2. Arquitectura del Sistema", 1)
    add_body(
        "El sistema sigue una arquitectura serverless basada en Azure Functions con "
        "integración a múltiples servicios cognitivos y de datos de Azure."
    )

    add_heading("2.1 Componentes Principales", 2, C_BLUE)
    components = [
        ("Dynamics 365 CRM",    "Origen de oportunidades comerciales"),
        ("Power Automate",      "Orquestador de flujo – dispara el análisis vía HTTP POST"),
        ("Azure Functions v4",  "Motor de procesamiento serverless (func-analyzer-prod)"),
        ("Azure OpenAI",        "GPT-4o-mini para análisis inteligente"),
        ("Azure AI Search",     "Base de conocimiento de torres y equipos"),
        ("Azure Blob Storage",  "Almacenamiento de PDFs generados"),
        ("Azure Cosmos DB",     "Historial y auditoría de análisis"),
        ("Microsoft Teams",     "Destino de la Adaptive Card con resultados"),
    ]
    styled_table(
        ["Componente", "Rol"],
        components,
        col_widths=[5.5, 11],
    )

    add_heading("2.2 Estructura de Código", 2, C_BLUE)
    structure = [
        ("AnalyzeOpportunity/__init__.py", "Entrypoint HTTP de la Function"),
        ("AnalyzeOpportunity/function.json", "Configuración del trigger (HTTP, auth level)"),
        ("shared/core/orchestrator.py",   "Orquestador principal del flujo de análisis"),
        ("shared/services/openai_service.py", "Integración con Azure OpenAI"),
        ("shared/services/search_service.py", "Integración con Azure AI Search"),
        ("shared/services/blob_storage_service.py", "Integración con Blob Storage"),
        ("shared/services/cosmos_service.py", "Integración con Cosmos DB"),
        ("shared/generators/pdf_generator.py", "Generación de PDFs con ReportLab"),
        ("shared/generators/adaptive_card.py", "Construcción de Adaptive Cards"),
        ("shared/models/opportunity.py",  "Modelo Pydantic de validación del payload"),
        ("data/torres_data_prod.json",    "Datos de respaldo de equipos"),
        ("host.json",                     "Configuración de la Function App"),
    ]
    styled_table(
        ["Archivo", "Descripción"],
        structure,
        col_widths=[7, 9.5],
    )

    # ── 3. FLUJO DE PROCESAMIENTO ────────────────────────────
    doc.add_paragraph()
    add_heading("3. Flujo de Procesamiento", 1)
    for step, title, desc in DocData.FLOW_STEPS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        rb = p.add_run(f"  Paso {step}:  ")
        rb.bold = True
        rb.font.color.rgb = rgb(*C_ACCENT)
        rb.font.size = Pt(10)
        rt = p.add_run(title)
        rt.bold = True
        rt.font.size = Pt(10)
        rt.font.color.rgb = rgb(*C_NAVY)
        add_body(desc, indent=True)

    # ── 4. API REFERENCE ────────────────────────────────────
    doc.add_paragraph()
    add_heading("4. API Reference", 1)

    add_heading("4.1 Endpoint", 2, C_BLUE)
    add_code("POST https://func-analyzer-prod.azurewebsites.net/api/analyze?code=<FUNCTION_KEY>")
    add_kv("Content-Type",    "application/json")
    add_kv("Autenticación",   "Function Key (query param 'code')")
    add_kv("Timeout recomendado", "120 segundos (el análisis IA toma ~15-30 s)")

    add_heading("4.2 Payload de Entrada", 2, C_BLUE)
    add_body("Todos los campos string; los marcados ✅ son requeridos.")
    styled_table(
        ["Campo", "Tipo", "Req.", "Descripción"],
        DocData.PAYLOAD_FIELDS,
        col_widths=[5, 1.8, 1.2, 8.5],
    )

    add_heading("4.3 Ejemplo de Payload", 2, C_BLUE)
    example_payload = '''{
  "opportunityid": "3e34c311-9c1f-448a-9c9e-6a680f28ccaa",
  "name": "Gobierno de Datos",
  "description": "Implementación de herramienta de Gobierno de Datos.",
  "estimatedvalue": 50000,
  "customerid": "cliente-id-001",
  "customername": "Empresa Ejemplo S.A.",
  "sdk_message": "Create"
}'''
    for line in example_payload.split("\n"):
        add_code(line)

    add_heading("4.4 Estructura de Respuesta Exitosa", 2, C_BLUE)
    response_fields = [
        ("success",                     "boolean", "true si el análisis se completó"),
        ("opportunity_id",              "string",  "ID de la oportunidad procesada"),
        ("opportunity_name",            "string",  "Nombre de la oportunidad"),
        ("event_type",                  "string",  "Create / Update / Unknown"),
        ("analysis.executive_summary",  "string",  "Resumen ejecutivo del análisis"),
        ("analysis.required_towers",    "array",   "Nombres de torres recomendadas"),
        ("analysis.team_recommendations","array",  "Detalle de equipos con líder y score"),
        ("analysis.risks",              "array",   "Riesgos identificados con mitigación"),
        ("analysis.timeline_estimate",  "object",  "Duración y fases del proyecto"),
        ("analysis.effort_estimate",    "object",  "Horas, complejidad y tamaño del equipo"),
        ("analysis.confidence",         "number",  "Nivel de confianza del análisis (0-1)"),
        ("outputs.adaptive_card",       "object",  "Adaptive Card JSON para Teams v1.4"),
        ("outputs.pdf_url",             "string",  "URL SAS al PDF en Blob Storage (90 días)"),
        ("outputs.cosmos_record_id",    "string",  "ID del registro en Cosmos DB"),
        ("metadata.processing_time_seconds","number","Tiempo total de procesamiento"),
        ("metadata.model_used",         "string",  "Modelo de IA utilizado"),
        ("metadata.teams_evaluated",    "integer", "Cantidad de equipos evaluados"),
    ]
    styled_table(
        ["Campo", "Tipo", "Descripción"],
        response_fields,
        col_widths=[6.5, 2, 8],
    )

    add_heading("4.5 Códigos de Error", 2, C_BLUE)
    styled_table(
        ["Código", "HTTP", "Descripción"],
        DocData.ERROR_CODES,
        col_widths=[4.5, 1.5, 10.5],
    )

    # ── 5. SERVICIOS ──────────────────────────────────────────
    doc.add_paragraph()
    add_heading("5. Servicios y Dependencias", 1)
    for i, svc in enumerate(DocData.SERVICES):
        add_heading(f"5.{i+1} {svc['name']}", 2, C_BLUE)
        add_body(svc["description"])
        add_kv("Archivo",        svc["file"])
        add_kv("Método clave",   svc["key_method"])
        p = doc.add_paragraph()
        rb = p.add_run("Variables de entorno requeridas:  ")
        rb.bold = True
        rb.font.size = Pt(10)
        for ev in svc["env_vars"]:
            add_code(ev)

    # ── 6. CONFIGURACIÓN (APP SETTINGS) ──────────────────────
    doc.add_paragraph()
    add_heading("6. Configuración – Application Settings", 1)
    add_body(
        "Las variables de entorno se configuran en Azure Portal → "
        "func-analyzer-prod → Configuration → Application settings."
    )
    styled_table(
        ["Variable", "Req.", "Descripción", "Ejemplo"],
        DocData.APP_SETTINGS,
        col_widths=[5.5, 1.2, 5.8, 4],
    )

    # ── 7. CATÁLOGO DE TORRES ────────────────────────────────
    doc.add_paragraph()
    add_heading("7. Catálogo de Torres y Equipos", 1)
    add_body(
        f"La base de conocimiento contiene {len(torres)} torres especializadas "
        "indexadas en Azure AI Search. Cada torre incluye habilidades, tecnologías, "
        "frameworks y datos de contacto del líder."
    )

    rows_torres = [
        (t["id"], t["team_name"], t["tower"],
         t["team_lead"], t["team_lead_email"])
        for t in torres
    ]
    styled_table(
        ["#", "Equipo", "Torre", "Líder", "Email"],
        rows_torres,
        col_widths=[0.7, 3.2, 4.5, 4, 4.1],
    )

    # Detalle habilidades
    add_heading("7.1 Competencias por Torre", 2, C_BLUE)
    for t in torres:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        rb = p.add_run(f"  {t['team_name']}  ")
        rb.bold = True
        rb.font.size = Pt(10)
        rb.font.color.rgb = rgb(*C_BLUE)
        p.add_run(f"— {t['description']}")

        skills_text = " • ".join(t.get("skills", [])[:6])
        add_body(f"Skills: {skills_text}", indent=True)
        tech_text = ", ".join(t.get("technologies", [])[:8])
        add_body(f"Tecnologías: {tech_text}", indent=True)

    # ── 8. CI/CD ──────────────────────────────────────────────
    doc.add_paragraph()
    add_heading("8. Pipeline CI/CD – GitHub Actions", 1)
    add_body(
        f"El repositorio {DocData.REPO} dispone de un workflow automático que "
        "se activa en cada push a la rama master. El flujo garantiza calidad del "
        "código antes de cada despliegue."
    )

    styled_table(
        ["Paso", "Herramienta / Comando", "Descripción"],
        [(s[0], s[1], s[2]) for s in DocData.CICD_STEPS],
        col_widths=[2.5, 5, 9],
    )

    add_heading("8.1 Secreto requerido en GitHub", 2, C_BLUE)
    add_kv("Secret Name", "AZURE_FUNCTIONAPP_PUBLISH_PROFILE")
    add_body(
        "Obtener desde: Azure Portal → func-analyzer-prod → "
        "Deployment Center → Manage publish profile → Download."
    )

    # ── 9. SEGURIDAD ──────────────────────────────────────────
    doc.add_paragraph()
    add_heading("9. Consideraciones de Seguridad", 1)
    security_items = [
        "La Function Key se pasa como query param 'code'. No exponer en logs públicos.",
        "Las App Settings con secretos (OpenAI Key, Search Key, etc.) están cifradas en reposo en Azure.",
        "Las URLs SAS de PDFs tienen vigencia de 90 días con permisos de solo lectura.",
        "El archivo local.settings.json contiene placeholders; los secretos reales van SOLO en Azure Portal.",
        "No se almacenan datos personales del cliente en el PDF ni en la Adaptive Card.",
        "Cosmos DB persiste únicamente datos estructurados del análisis, no documentos del CRM.",
        "El timeout de 10 minutos evita costos excesivos por ejecuciones colgadas.",
    ]
    for item in security_items:
        add_bullet(item)

    # ── 10. PRUEBAS ───────────────────────────────────────────
    doc.add_paragraph()
    add_heading("10. Suite de Pruebas", 1)
    add_kv("Comando",         "python -m pytest tests/ -v")
    add_kv("Tests",           "22 pruebas unitarias")
    add_kv("Cobertura",       "Modelos Pydantic, validación de payload, enriquecimiento de equipos")
    add_kv("Linting",         "flake8 – max-line-length=120")
    doc.add_paragraph()
    add_body(
        "Los tests se ejecutan automáticamente en el pipeline CI/CD antes de cada "
        "despliegue. Un fallo en los tests bloquea el deploy."
    )

    # ── GUARDAR ───────────────────────────────────────────────
    doc.save(output_path)
    print(f"  ✅ DOCX guardado: {output_path}")


# ══════════════════════════════════════════════════════════════
#  GENERADOR  PDF
# ══════════════════════════════════════════════════════════════
def generate_pdf(output_path: str, torres: list) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, KeepTogether,
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

    # ── Colores RL ──
    NAVY    = colors.Color(*[x/255 for x in C_NAVY])
    BLUE    = colors.Color(*[x/255 for x in C_BLUE])
    ACCENT  = colors.Color(*[x/255 for x in C_ACCENT])
    ORANGE  = colors.Color(*[x/255 for x in C_ORANGE])
    LGRAY   = colors.Color(*[x/255 for x in C_LIGHT_GRAY])
    MGRAY   = colors.Color(*[x/255 for x in C_MID_GRAY])
    DGRAY   = colors.Color(*[x/255 for x in C_DARK_GRAY])
    WHITE   = colors.white
    SUCCESS = colors.Color(*[x/255 for x in C_SUCCESS])

    W, H = A4
    MARGIN = 2.2 * cm

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=MARGIN,  bottomMargin=MARGIN,
        title=DocData.PROJECT,
        author=DocData.COMPANY,
        subject=DocData.SUBTITLE,
    )

    # ── Estilos ──
    styles = getSampleStyleSheet()

    def style(name, parent="Normal", **kw):
        s = ParagraphStyle(name, parent=styles[parent], **kw)
        return s

    S = {
        "cover_title": style("ct", fontSize=24, textColor=NAVY,
                             fontName="Helvetica-Bold", alignment=TA_CENTER,
                             spaceAfter=12),
        "cover_sub":   style("cs", fontSize=14, textColor=BLUE,
                             fontName="Helvetica", alignment=TA_CENTER,
                             spaceAfter=8),
        "cover_info":  style("ci", fontSize=11, textColor=DGRAY,
                             alignment=TA_CENTER, spaceAfter=4),
        "h1":          style("h1", fontSize=14, textColor=NAVY,
                             fontName="Helvetica-Bold", spaceBefore=14,
                             spaceAfter=4),
        "h2":          style("h2", fontSize=11, textColor=BLUE,
                             fontName="Helvetica-Bold", spaceBefore=8,
                             spaceAfter=3),
        "h3":          style("h3", fontSize=10, textColor=ACCENT,
                             fontName="Helvetica-BoldOblique", spaceBefore=5,
                             spaceAfter=2),
        "body":        style("body", fontSize=9.5, textColor=DGRAY,
                             leading=14, alignment=TA_JUSTIFY, spaceAfter=4),
        "bullet":      style("blt", fontSize=9.5, textColor=DGRAY,
                             leading=13, leftIndent=14, bulletIndent=4,
                             spaceAfter=2),
        "code":        style("code", fontSize=8.5, textColor=colors.Color(0.7, 0.1, 0.1),
                             fontName="Courier", leftIndent=16, spaceAfter=2),
        "kv_key":      style("kvk", fontSize=9.5, textColor=NAVY,
                             fontName="Helvetica-Bold", spaceAfter=2),
        "tbl_hdr":     style("thdr", fontSize=8.5, textColor=WHITE,
                             fontName="Helvetica-Bold", alignment=TA_CENTER),
        "tbl_cell":    style("tcll", fontSize=8.5, textColor=DGRAY,
                             leading=12),
        "footer":      style("ftr", fontSize=7.5, textColor=MGRAY,
                             alignment=TA_CENTER),
    }

    # ── Header / Footer ──
    def on_page(canvas, doc_):
        canvas.saveState()
        # Header bar
        canvas.setFillColor(NAVY)
        canvas.rect(0, H - 1.2 * cm, W, 1.2 * cm, fill=1, stroke=0)
        canvas.setFillColor(WHITE)
        canvas.setFont("Helvetica-Bold", 8)
        canvas.drawString(MARGIN, H - 0.75 * cm, DocData.PROJECT)
        canvas.drawRightString(W - MARGIN, H - 0.75 * cm, DocData.VERSION)
        # Footer bar
        canvas.setFillColor(LGRAY)
        canvas.rect(0, 0, W, 1.0 * cm, fill=1, stroke=0)
        canvas.setFillColor(DGRAY)
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(MARGIN, 0.35 * cm, DocData.COMPANY)
        canvas.drawRightString(
            W - MARGIN, 0.35 * cm, f"Página {doc_.page}"
        )
        canvas.restoreState()

    def on_first_page(canvas, doc_):
        canvas.saveState()
        # Franja superior decorativa
        canvas.setFillColor(NAVY)
        canvas.rect(0, H - 4 * cm, W, 4 * cm, fill=1, stroke=0)
        canvas.setFillColor(ACCENT)
        canvas.rect(0, H - 4.3 * cm, W, 0.3 * cm, fill=1, stroke=0)
        # Footer
        canvas.setFillColor(LGRAY)
        canvas.rect(0, 0, W, 1.0 * cm, fill=1, stroke=0)
        canvas.setFillColor(DGRAY)
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(MARGIN, 0.35 * cm, DocData.COMPANY)
        canvas.drawRightString(W - MARGIN, 0.35 * cm, DocData.DATE)
        canvas.restoreState()

    # ── Helpers ──
    def h1(text):       return [Spacer(1, 4), Paragraph(text, S["h1"]),
                                 HRFlowable(width="100%", thickness=1.5,
                                            color=ACCENT, spaceAfter=4)]
    def h2(text):       return [Paragraph(text, S["h2"])]
    def h3(text):       return [Paragraph(text, S["h3"])]
    def body(text):     return [Paragraph(text, S["body"])]
    def bullet(text):   return [Paragraph(f"• {text}", S["bullet"])]
    def code(text):     return [Paragraph(text, S["code"])]
    def sp(h=0.25):     return [Spacer(1, h * cm)]
    def kv(key, val):
        return [Paragraph(f"<b>{key}:</b>  {val}", S["body"])]

    def mk_table(headers, rows, col_w, h_color=NAVY):
        data = [[Paragraph(h, S["tbl_hdr"]) for h in headers]]
        for ri, row in enumerate(rows):
            bg = LGRAY if ri % 2 == 0 else WHITE
            data.append([Paragraph(str(c), S["tbl_cell"]) for c in row])
        t = Table(data, colWidths=[w * cm for w in col_w])
        ts = TableStyle([
            ("BACKGROUND",  (0, 0), (-1, 0),  h_color),
            ("TEXTCOLOR",   (0, 0), (-1, 0),  WHITE),
            ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0, 0), (-1, 0),  8),
            ("ALIGN",       (0, 0), (-1, 0),  "CENTER"),
            ("GRID",        (0, 0), (-1, -1), 0.4, MGRAY),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LGRAY, WHITE]),
            ("FONTSIZE",    (0, 1), (-1, -1), 8),
            ("VALIGN",      (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",  (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ])
        t.setStyle(ts)
        return [t, Spacer(1, 0.3 * cm)]

    # ── CONTENIDO ──────────────────────────────────────────────
    story = []
    content_w = W - 2 * MARGIN

    # Portada
    story += [Spacer(1, 3.8 * cm)]
    story += [Paragraph(f'<font color="#FFFFFF"><b>{DocData.PROJECT}</b></font>', S["cover_title"])]
    story += [Paragraph(f'<font color="#FFFFFF">{DocData.SUBTITLE}</font>', S["cover_sub"])]
    story += [Spacer(1, 1.5 * cm)]
    for label, val in [("Versión", DocData.VERSION), ("Fecha", DocData.DATE),
                       ("Empresa", DocData.COMPANY)]:
        story += [Paragraph(f"<b>{label}:</b>  {val}", S["cover_info"])]
    story += [PageBreak()]

    # 1. Resumen ejecutivo
    story += h1("1. Resumen Ejecutivo")
    story += body(DocData.SUMMARY)
    story += sp()
    for k, v in [
        ("Function App", DocData.AZURE_FUNC),
        ("Resource Group", DocData.AZURE_RG),
        ("Endpoint", DocData.FUNCTION_URL),
        ("Modelo IA", "GPT-4o-mini (Azure OpenAI)"),
        ("Runtime", "Python 3.13 – Azure Functions v4"),
        ("Timeout", "10 minutos"),
        ("CI/CD", "GitHub Actions → Publish Profile"),
    ]:
        story += kv(k, v)

    # 2. Arquitectura
    story += h1("2. Arquitectura del Sistema")
    story += body(
        "El sistema sigue una arquitectura serverless basada en Azure Functions "
        "con integración a múltiples servicios cognitivos y de datos de Azure."
    )
    story += h2("2.1 Componentes Principales")
    story += mk_table(
        ["Componente", "Rol"],
        [("Dynamics 365 CRM",   "Origen de oportunidades comerciales"),
         ("Power Automate",     "Orquestador de flujo – dispara el análisis vía HTTP POST"),
         ("Azure Functions v4", "Motor de procesamiento serverless (func-analyzer-prod)"),
         ("Azure OpenAI",       "GPT-4o-mini para análisis inteligente"),
         ("Azure AI Search",    "Base de conocimiento de torres y equipos"),
         ("Azure Blob Storage", "Almacenamiento de PDFs generados"),
         ("Azure Cosmos DB",    "Historial y auditoría de análisis"),
         ("Microsoft Teams",    "Destino de la Adaptive Card con resultados")],
        [5.5, 11.3],
    )
    story += h2("2.2 Estructura de Código")
    story += mk_table(
        ["Archivo", "Descripción"],
        [("AnalyzeOpportunity/__init__.py",        "Entrypoint HTTP de la Function"),
         ("shared/core/orchestrator.py",            "Orquestador principal del flujo"),
         ("shared/services/openai_service.py",      "Integración con Azure OpenAI"),
         ("shared/services/search_service.py",      "Integración con Azure AI Search"),
         ("shared/services/blob_storage_service.py","Integración con Blob Storage"),
         ("shared/services/cosmos_service.py",      "Integración con Cosmos DB"),
         ("shared/generators/pdf_generator.py",     "Generación de PDFs con ReportLab"),
         ("shared/generators/adaptive_card.py",     "Construcción de Adaptive Cards"),
         ("shared/models/opportunity.py",           "Modelo Pydantic – validación payload"),
         ("host.json",                              "Configuración de la Function App")],
        [7, 9.8],
    )

    # 3. Flujo
    story += h1("3. Flujo de Procesamiento")
    for step, title, desc in DocData.FLOW_STEPS:
        story += [Paragraph(
            f'<b><font color="#{C_ACCENT[0]:02X}{C_ACCENT[1]:02X}{C_ACCENT[2]:02X}">'
            f'Paso {step}</font></b>  –  <b>{title}</b>', S["body"]
        )]
        story += body(f"    {desc}")

    # 4. API Reference
    story += h1("4. API Reference")
    story += h2("4.1 Endpoint")
    story += code("POST https://func-analyzer-prod.azurewebsites.net/api/analyze?code=&lt;FUNCTION_KEY&gt;")
    story += kv("Content-Type", "application/json")
    story += kv("Autenticación", "Function Key (query param 'code')")
    story += kv("Timeout recomendado", "120 segundos")

    story += h2("4.2 Payload de Entrada")
    story += mk_table(
        ["Campo", "Tipo", "Req.", "Descripción"],
        DocData.PAYLOAD_FIELDS,
        [4.5, 1.8, 1.0, 9.5],
    )

    story += h2("4.3 Estructura de Respuesta")
    story += mk_table(
        ["Campo", "Tipo", "Descripción"],
        [
            ("success",                      "boolean", "true si el análisis se completó"),
            ("analysis.executive_summary",   "string",  "Resumen ejecutivo del análisis"),
            ("analysis.required_towers",     "array",   "Torres recomendadas"),
            ("analysis.team_recommendations","array",   "Equipos con líder, score y justificación"),
            ("analysis.risks",               "array",   "Riesgos con probabilidad y mitigación"),
            ("analysis.timeline_estimate",   "object",  "Duración y fases"),
            ("analysis.effort_estimate",     "object",  "Horas, complejidad, tamaño equipo"),
            ("analysis.confidence",          "number",  "Confianza del análisis (0-1)"),
            ("outputs.adaptive_card",        "object",  "Adaptive Card JSON para Teams"),
            ("outputs.pdf_url",              "string",  "URL SAS al PDF (90 días)"),
            ("outputs.cosmos_record_id",     "string",  "ID del registro en Cosmos DB"),
            ("metadata.processing_time_seconds","number","Tiempo de procesamiento"),
        ],
        [5.5, 2, 9.3],
    )

    story += h2("4.4 Códigos de Error")
    story += mk_table(
        ["Código", "HTTP", "Descripción"],
        DocData.ERROR_CODES,
        [4.5, 1.2, 11.1],
    )

    # 5. Servicios
    story += h1("5. Servicios y Dependencias")
    for i, svc in enumerate(DocData.SERVICES):
        story += h2(f"5.{i+1} {svc['name']}")
        story += body(svc["description"])
        story += kv("Archivo", svc["file"])
        story += kv("Método clave", svc["key_method"])
        vars_str = "  |  ".join(svc["env_vars"])
        story += body(f"<b>Variables requeridas:</b>  {vars_str}")
        story += sp(0.2)

    # 6. App Settings
    story += h1("6. Configuración – Application Settings")
    story += body(
        "Variables de entorno que deben configurarse en Azure Portal → "
        "func-analyzer-prod → Configuration → Application settings."
    )
    story += mk_table(
        ["Variable", "Req.", "Descripción", "Ejemplo"],
        DocData.APP_SETTINGS,
        [5.2, 1.0, 5.5, 5.1],
    )

    # 7. Torres
    story += h1("7. Catálogo de Torres y Equipos")
    story += body(
        f"La base de conocimiento contiene {len(torres)} torres especializadas "
        "indexadas en Azure AI Search."
    )
    rows_t = [(t["id"], t["team_name"], t["tower"],
               t["team_lead"], t["team_lead_email"]) for t in torres]
    story += mk_table(
        ["#", "Equipo", "Torre", "Líder", "Email"],
        rows_t,
        [0.6, 2.8, 4.0, 3.8, 5.6],
    )
    story += h2("7.1 Competencias por Torre")
    for t in torres:
        sk = " • ".join(t.get("skills", [])[:5])
        te = ", ".join(t.get("technologies", [])[:6])
        story += [KeepTogether([
            Paragraph(f'<b><font color="#{C_BLUE[0]:02X}{C_BLUE[1]:02X}{C_BLUE[2]:02X}">'
                      f'{t["team_name"]}</font></b> — {t["description"]}', S["body"]),
            Paragraph(f'<i>Skills: {sk}</i>', S["body"]),
            Paragraph(f'<i>Tecnologías: {te}</i>', S["body"]),
            Spacer(1, 0.15 * cm),
        ])]

    # 8. CI/CD
    story += h1("8. Pipeline CI/CD – GitHub Actions")
    story += body(
        f"Cada push a la rama <b>master</b> de {DocData.REPO} "
        "dispara el pipeline automático de build, test y deploy."
    )
    story += mk_table(
        ["Paso", "Herramienta / Comando", "Descripción"],
        [(s[0], s[1], s[2]) for s in DocData.CICD_STEPS],
        [2.0, 5.5, 9.3],
    )
    story += h2("8.1 Secreto requerido en GitHub")
    story += kv("Secret Name", "AZURE_FUNCTIONAPP_PUBLISH_PROFILE")
    story += body(
        "Obtener desde: Azure Portal → func-analyzer-prod → "
        "Deployment Center → Manage publish profile → Download."
    )

    # 9. Seguridad
    story += h1("9. Consideraciones de Seguridad")
    for item in [
        "La Function Key se pasa como query param 'code'. No exponer en logs públicos.",
        "Las App Settings con secretos están cifradas en reposo en Azure.",
        "Las URLs SAS de PDFs tienen vigencia de 90 días con permisos de solo lectura.",
        "local.settings.json contiene placeholders. Los secretos reales van SOLO en Azure Portal.",
        "No se almacenan datos personales del cliente en el PDF ni en la Adaptive Card.",
        "Cosmos DB persiste únicamente datos estructurados del análisis.",
        "El timeout de 10 minutos evita costos excesivos por ejecuciones colgadas.",
    ]:
        story += bullet(item)

    # 10. Pruebas
    story += h1("10. Suite de Pruebas")
    story += kv("Comando",   "python -m pytest tests/ -v")
    story += kv("Tests",     "22 pruebas unitarias")
    story += kv("Cobertura", "Modelos Pydantic, validación de payload, enriquecimiento de equipos")
    story += kv("Linting",   "flake8 – max-line-length=120")
    story += body(
        "Los tests se ejecutan automáticamente en el pipeline CI/CD antes de cada despliegue. "
        "Un fallo en los tests bloquea el deploy a producción."
    )

    # Generar
    doc.build(story, onFirstPage=on_first_page, onLaterPages=on_page)
    print(f"  ✅ PDF guardado: {output_path}")


# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════
def main():
    base = Path(__file__).resolve().parent.parent
    data_file = base / "data" / "torres_data_prod.json"

    with open(data_file, encoding="utf-8") as f:
        torres = json.load(f)

    out_dir = base / "docs"
    out_dir.mkdir(exist_ok=True)

    today_str = date.today().strftime("%Y%m%d")
    docx_path = str(out_dir / f"Perxia_Solver_Documentacion_Tecnica_{today_str}.docx")
    pdf_path  = str(out_dir / f"Perxia_Solver_Documentacion_Tecnica_{today_str}.pdf")

    print("\n📄 Generando documentación técnica...\n")
    generate_docx(docx_path, torres)
    generate_pdf(pdf_path,  torres)
    print(f"\n✅ Documentación generada en: {out_dir}")


if __name__ == "__main__":
    main()
